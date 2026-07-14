import tempfile
from pathlib import Path

import pytest

from src.tools.file_readers.text_reader import TextReader
from src.tools.file_readers.pdf_reader import PDFReader
from src.tools.file_readers.docx_reader import DOCXReader
from src.tools.file_readers.markdown_reader import MarkdownReader
from src.tools.file_readers.python_reader import PythonReader
from src.tools.file_readers.excel_reader import ExcelReader
from src.tools.file_readers.image_reader import ImageReader
from src.tools.file_readers.router import FileRouter

class TestTextReader:
    def test_read_text(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("你好世界\nHello World")
            f.flush()

        try:
            reader = TextReader()
            content = reader.read(f.name)
            assert "你好世界" in content
            assert "Hello World" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

    def test_supported_extension(self):
        reader = TextReader()
        assert reader.can_handle(".txt")
        assert not reader.can_handle(".pdf")

class TestPDFReader:
    def test_read_pdf(self):
        try:
            from pypdf import PdfWriter
        except ImportError:
            pytest.skip("pypdf not installed")

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            writer = PdfWriter()
            writer.add_blank_page(width=100, height=100)
            writer.write(f)
            f.flush()

        try:
            reader = PDFReader()
            content = reader.read(f.name)
            assert isinstance(content, str)
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestDOCXReader:
    def test_read_docx(self):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("测试中文")
            doc.save(f.name)

        try:
            reader = DOCXReader()
            content = reader.read(f.name)
            assert "Hello World" in content
            assert "测试中文" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestMarkdownReader:
    def test_read_markdown(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write("# Title\n\nContent here.")
            f.flush()

        try:
            reader = MarkdownReader()
            content = reader.read(f.name)
            assert "# Title" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestPythonReader:
    def test_read_python(self):
        import tempfile
        f = tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False)
        f.write("import os\n\ndef main():\n    pass\n")
        f.flush()
        f.close()

        try:
            reader = PythonReader()
            content = reader.read(f.name)
            assert "def main" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestExcelReader:
    def test_read_excel(self):
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Sheet1"
            ws["A1"] = "Name"
            ws["B1"] = "Value"
            ws["A2"] = "Test"
            ws["B2"] = "123"
            wb.save(f.name)

        try:
            reader = ExcelReader()
            content = reader.read(f.name)
            assert "Sheet1" in content
            assert "Name" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestImageReader:
    def test_read_image(self):
        try:
            from PIL import Image
        except ImportError:
            pytest.skip("Pillow not installed")

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            img = Image.new("RGB", (100, 200), color="red")
            img.save(f, "PNG")

        try:
            reader = ImageReader()
            content = reader.read(f.name)
            assert "PNG" in content
            assert "100" in content
        finally:
            Path(f.name).unlink(missing_ok=True)

class TestFileRouter:
    def test_route_text_file(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("test content")
            f.flush()

        try:
            router = FileRouter()
            doc = router.route(f.name)
            assert doc.type == "txt"
            assert "test content" in doc.content
        finally:
            Path(f.name).unlink(missing_ok=True)

    def test_route_python_file(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False, encoding="utf-8") as f:
            f.write("print('hello')")
            f.flush()

        try:
            router = FileRouter()
            doc = router.route(f.name)
            assert doc.type == "py"
        finally:
            Path(f.name).unlink(missing_ok=True)

    def test_unsupported_extension(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False, encoding="utf-8") as f:
            f.write("data")
            f.flush()

        try:
            router = FileRouter()
            with pytest.raises(ValueError, match="Unsupported file type"):
                router.route(f.name)
        finally:
            Path(f.name).unlink(missing_ok=True)
